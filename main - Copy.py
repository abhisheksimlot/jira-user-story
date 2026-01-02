import os
import json
import io
import html
from typing import List, Dict, Any

from fastapi import FastAPI, Form
from fastapi.responses import HTMLResponse, StreamingResponse
from openai import OpenAI
from docx import Document
from docx.shared import Pt


# ------------- LLM FUNCTION (OpenAI) ------------- #

def generate_requirements_from_text(raw_text: str) -> Dict[str, Any]:
    """
    Call the OpenAI API to convert raw/conversational text into
    structured Jira-style requirements.
    Returns a Python dict matching our JSON schema.
    """

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError(
            "Missing OPENAI_API_KEY environment variable. "
            "Please set it before running this app."
        )

    client = OpenAI(api_key=api_key)

    system_prompt = """
You are a senior business analyst who writes clear, concise Jira-style requirements.

Given raw notes or conversation transcripts, extract a set of IMPLEMENTABLE requirements.

Rules:
- Identify each distinct feature or piece of functionality and generate a separate Jira Story for each.
- Do NOT merge multiple features into one story. One feature = one story.
- Use user story style when relevant:
  "As a <role>, I want <feature> so that <benefit>."
- Always include acceptance criteria in GIVEN-WHEN-THEN style where possible.
- If something is unclear, make a reasonable assumption and mention it in the description.
- Output VALID JSON following this exact structure:

{
  "project_key": "SHORT_PROJECT_CODE",
  "requirements": [
    {
      "id": "REQ-1",
      "issue_type": "Story" | "Task" | "Bug",
      "summary": "Short Jira summary line",
      "description": "Longer functional description with key details and assumptions.",
      "priority": "High" | "Medium" | "Low",
      "story_points": 1,
      "acceptance_criteria": [
        "GIVEN ... WHEN ... THEN ...",
        "..."
      ],
      "dependencies": [
        "REQ-2"
      ]
    }
  ]
}

- Use REQ-1, REQ-2, ... for IDs in order.
- Always include at least one acceptance_criteria item.
- If story_points is unclear, estimate a number between 1 and 13.
"""

    user_prompt = f"""
Source text (notes / conversation):

\"\"\"{raw_text}\"\"\"

Now produce the JSON only, with no extra commentary.
"""

    # Use JSON mode via the `text` parameter.
    # In JSON mode, the serialized JSON is returned in `response.output_text`. :contentReference[oaicite:1]{index=1}
    response = client.responses.create(
        model="gpt-4o-mini",  # or another model you have access to
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        max_output_tokens=2000,
        text={"format": {"type": "json_object"}},
    )

    # The whole JSON string is here:
    json_text = response.output_text

    if not json_text or not json_text.strip():
        # This is what was causing "Expecting value: line 1 column 1 (char 0)"
        raise ValueError("The model returned an empty response. Please try again.")

    try:
        data = json.loads(json_text)
    except json.JSONDecodeError as e:
        # Give a clearer error back to the web page
        preview = json_text[:200].replace("\n", " ")
        raise ValueError(
            f"Model output was not valid JSON. "
            f"Error: {str(e)}. Output preview: {preview}"
        )

    return data



# ------------- WORD DOC CREATION (in memory) ------------- #

def create_word_bytes_from_requirements(data: Dict[str, Any]) -> bytes:
    """
    Create a Word document from the structured requirements JSON using python-docx.
    Return the document as bytes (so we can send it as a download).
    """

    project_key = data.get("project_key", "PROJECT")
    requirements: List[Dict[str, Any]] = data.get("requirements", [])

    doc = Document()

    # Title
    title = doc.add_heading(level=0)
    title_run = title.add_run(f"{project_key} – Requirements Specification")
    title_run.font.size = Pt(20)

    doc.add_paragraph()  # blank line

    # Overview
    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "This document contains auto-generated Jira-style requirements derived "
        "from raw text and conversational input."
    )

    doc.add_page_break()

    # Requirements
    for req in requirements:
        req_id = req.get("id", "REQ-?")
        issue_type = req.get("issue_type", "Story")
        summary = req.get("summary", "")
        description = req.get("description", "")
        priority = req.get("priority", "Medium")
        story_points = req.get("story_points", None)
        acceptance_criteria = req.get("acceptance_criteria", [])
        dependencies = req.get("dependencies", [])

        # Requirement header
        doc.add_heading(f"{req_id}: {summary}", level=1)

        # Meta info
        meta_para = doc.add_paragraph()
        meta_para.add_run("Issue Type: ").bold = True
        meta_para.add_run(issue_type)

        meta_para.add_run("   |   Priority: ").bold = True
        meta_para.add_run(priority)

        if story_points is not None:
            meta_para.add_run("   |   Story Points: ").bold = True
            meta_para.add_run(str(story_points))

        if dependencies:
            meta_para = doc.add_paragraph()
            meta_para.add_run("Dependencies: ").bold = True
            meta_para.add_run(", ".join(dependencies))

        # Description
        doc.add_heading("Description", level=2)
        doc.add_paragraph(description)

        # Acceptance Criteria
        doc.add_heading("Acceptance Criteria", level=2)
        for criterion in acceptance_criteria:
            doc.add_paragraph(criterion, style="List Bullet")

        doc.add_page_break()

    # Save to bytes instead of file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ------------- HELPER: RENDER FORM WITH OPTIONAL ERROR ------------- #

def render_form_page(
    conversation_text: str = "",
    project_key: str = "",
    error_message: str | None = None,
) -> HTMLResponse:
    """
    Render the HTML form. If error_message is provided, show it at the top.
    Also pre-fill conversation_text and project_key so the user doesn't lose input.
    """

    error_html = ""
    if error_message:
        error_html = f"""
        <div class="error">
            {html.escape(error_message)}
        </div>
        """

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Conversation to Jira Stories</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                margin: 40px;
                max-width: 800px;
            }}
            textarea {{
                width: 100%;
                height: 300px;
                font-family: Consolas, monospace;
                font-size: 14px;
            }}
            label {{
                font-weight: bold;
            }}
            button {{
                margin-top: 20px;
                padding: 10px 20px;
                font-size: 16px;
            }}
            .container {{
                display: flex;
                flex-direction: column;
                gap: 10px;
            }}
            .error {{
                background-color: #ffe5e5;
                border: 1px solid #ff8080;
                color: #b00020;
                padding: 10px;
                border-radius: 4px;
                margin-bottom: 15px;
            }}
        </style>
    </head>
    <body>
        <h1>Conversation → Jira Stories (Word)</h1>
        <p>Paste your meeting notes or conversation text below and click <strong>Generate Word</strong>.</p>
        {error_html}
        <form method="post" action="/generate">
            <div class="container">
                <label for="conversation_text">Conversation / Notes:</label>
                <textarea id="conversation_text" name="conversation_text" required>{html.escape(conversation_text)}</textarea>
                <label for="project_key">Project Key (optional, e.g. CLAIMS):</label>
                <input type="text" id="project_key" name="project_key" placeholder="CLAIMS" value="{html.escape(project_key)}" />
                <button type="submit">Generate Word</button>
            </div>
        </form>
    </body>
    </html>
    """
    return HTMLResponse(content=html_content)


# ------------- FASTAPI APP ------------- #

app = FastAPI()


# Simple HTML form page
@app.get("/", response_class=HTMLResponse)
async def form_page():
    # No error by default
    return render_form_page()


# Handle form submit and return Word file
@app.post("/generate")
async def generate_word(conversation_text: str = Form(...), project_key: str = Form(None)):
    try:
        # 1) Call OpenAI to generate requirements
        data = generate_requirements_from_text(conversation_text)

        # 2) Override project key if user provided one
        if project_key:
            data["project_key"] = project_key

        # 3) Create Word document bytes
        doc_bytes = create_word_bytes_from_requirements(data)

        # 4) Stream as a downloadable file
        filename = f"{data.get('project_key', 'PROJECT')}_requirements.docx"

        return StreamingResponse(
            io.BytesIO(doc_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            },
        )

    except ValueError as ve:
        # Typically for missing OPENAI_API_KEY or validation-type issues
        return render_form_page(
            conversation_text=conversation_text,
            project_key=project_key or "",
            error_message=str(ve),
        )

    except Exception as e:
        # Generic fallback
        return render_form_page(
            conversation_text=conversation_text,
            project_key=project_key or "",
            error_message="Something went wrong while generating the Word file. "
                          "Please try again or contact support. "
                          f"Details: {str(e)}",
        )
