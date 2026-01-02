import os
import json
import io
import html
from typing import List, Dict, Any

from fastapi import FastAPI, Form, UploadFile, File
from fastapi.responses import HTMLResponse, StreamingResponse
from openai import OpenAI
from docx import Document
from docx.shared import Pt


# ------------- LLM FUNCTION (OpenAI) ------------- #

def generate_requirements_from_text(raw_text: str) -> Dict[str, Any]:
    """
    Call the OpenAI API to convert raw/conversational text into
    structured Jira-style requirements.
    Returns a Python dict matching our JSON schema.
    """

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError(
            "Missing OPENAI_API_KEY environment variable. "
            "Please set it before running this app."
        )

    client = OpenAI(api_key=api_key)

    system_prompt = """
You are a senior business analyst who writes clear, concise Jira-style requirements.

Given raw notes or conversation transcripts, extract a set of IMPLEMENTABLE requirements.

Rules:
- Identify each distinct feature or piece of functionality and generate a separate Jira Story for each.
- Do NOT merge multiple features into one story. One feature = one story.
- Use user story style when relevant:
  "As a <role>, I want <feature> so that <benefit>."
- Always include acceptance criteria in GIVEN-WHEN-THEN style where possible.
- If something is unclear, make a reasonable assumption and mention it in the description.
- Output VALID JSON following this exact structure:

{
  "project_key": "SHORT_PROJECT_CODE",
  "requirements": [
    {
      "id": "REQ-1",
      "issue_type": "Story" | "Task" | "Bug",
      "summary": "Short Jira summary line",
      "description": "Longer functional description with key details and assumptions.",
      "priority": "High" | "Medium" | "Low",
      "story_points": 1,
      "acceptance_criteria": [
        "GIVEN ... WHEN ... THEN ...",
        "..."
      ],
      "dependencies": [
        "REQ-2"
      ]
    }
  ]
}

- Use REQ-1, REQ-2, ... for IDs in order.
- Always include at least one acceptance_criteria item.
- If story_points is unclear, estimate a number between 1 and 13.
"""

    user_prompt = f"""
Source text (notes / conversation):

\"\"\"{raw_text}\"\"\"

Now produce the JSON only, with no extra commentary.
"""

    # Use JSON mode via the `text` parameter.
    # In JSON mode, the serialized JSON is returned in `response.output_text`. :contentReference[oaicite:1]{index=1}
    response = client.responses.create(
        model="gpt-4o-mini",  # or another model you have access to
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        max_output_tokens=2000,
        text={"format": {"type": "json_object"}},
    )

    # The whole JSON string is here:
    json_text = response.output_text

    if not json_text or not json_text.strip():
        # This is what was causing "Expecting value: line 1 column 1 (char 0)"
        raise ValueError("The model returned an empty response. Please try again.")

    try:
        data = json.loads(json_text)
    except json.JSONDecodeError as e:
        # Give a clearer error back to the web page
        preview = json_text[:200].replace("\n", " ")
        raise ValueError(
            f"Model output was not valid JSON. "
            f"Error: {str(e)}. Output preview: {preview}"
        )

    return data



# ------------- WORD DOC CREATION (in memory) ------------- #

def create_word_bytes_from_requirements(data: Dict[str, Any]) -> bytes:
    """
    Create a Word document from the structured requirements JSON using python-docx.
    Return the document as bytes (so we can send it as a download).
    """

    project_key = data.get("project_key", "PROJECT")
    requirements: List[Dict[str, Any]] = data.get("requirements", [])

    doc = Document()

    # Title
    title = doc.add_heading(level=0)
    title_run = title.add_run(f"{project_key} – Requirements Specification")
    title_run.font.size = Pt(20)

    doc.add_paragraph()  # blank line

    # Overview
    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "This document contains auto-generated Jira-style requirements derived "
        "from raw text and conversational input."
    )

    doc.add_page_break()

    # Requirements
    for req in requirements:
        req_id = req.get("id", "REQ-?")
        issue_type = req.get("issue_type", "Story")
        summary = req.get("summary", "")
        description = req.get("description", "")
        priority = req.get("priority", "Medium")
        story_points = req.get("story_points", None)
        acceptance_criteria = req.get("acceptance_criteria", [])
        dependencies = req.get("dependencies", [])

        # Requirement header
        doc.add_heading(f"{req_id}: {summary}", level=1)

        # Meta info
        meta_para = doc.add_paragraph()
        meta_para.add_run("Issue Type: ").bold = True
        meta_para.add_run(issue_type)

        meta_para.add_run("   |   Priority: ").bold = True
        meta_para.add_run(priority)

        if story_points is not None:
            meta_para.add_run("   |   Story Points: ").bold = True
            meta_para.add_run(str(story_points))

        if dependencies:
            meta_para = doc.add_paragraph()
            meta_para.add_run("Dependencies: ").bold = True
            meta_para.add_run(", ".join(dependencies))

        # Description
        doc.add_heading("Description", level=2)
        doc.add_paragraph(description)

        # Acceptance Criteria
        doc.add_heading("Acceptance Criteria", level=2)
        for criterion in acceptance_criteria:
            doc.add_paragraph(criterion, style="List Bullet")

        doc.add_page_break()

    # Save to bytes instead of file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ------------- HELPER: RENDER FORM WITH OPTIONAL ERROR ------------- #

def render_form_page(
    conversation_text: str = "",
    project_key: str = "",
    input_mode: str = "text",
    error_message: str | None = None,
) -> HTMLResponse:

    error_html = ""
    if error_message:
        error_html = f"""
        <div class="alert-error">
            <span>⚠️ {html.escape(error_message)}</span>
        </div>
        """

    text_checked = "checked" if input_mode == "text" else ""
    file_checked = "checked" if input_mode == "file" else ""

    # Display logic based on selected mode
    text_display = "flex" if input_mode == "text" else "none"
    file_display = "flex" if input_mode == "file" else "none"

    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <title>Conversation → Jira Stories</title>
        <meta charset="utf-8" />
        <meta name="viewport" content="width=device-width, initial-scale=1" />
        <style>
            :root {{
                --primary: #D32F2F;
                --primary-dark: #B71C1C;
                --primary-light: #FFCDD2;
                --bg-light: #ffffff;
                --card-shadow: rgba(211, 47, 47, 0.25);
                --border: #e5e7eb;
                --text-main: #1a1a1a;
                --text-muted: #6b6b6b;
                --error-bg: #ffebee;
                --error-border: #ef9a9a;
                --error-text: #c62828;
            }}

            body {{
                margin: 0;
                font-family: 'Segoe UI', Arial, sans-serif;
                background: linear-gradient(120deg, #ffffff, #ffebee, #ffffff);
                display: flex;
                justify-content: center;
                align-items: center;
                padding: 30px;
                min-height: 100vh;
            }}

            .app-shell {{
                width: 100%;
                max-width: 900px;
            }}

            .card {{
                background: #fff;
                border-radius: 18px;
                padding: 30px;
                box-shadow: 0 15px 35px var(--card-shadow);
                border-left: 6px solid var(--primary);
                display: flex;
                flex-direction: column;
                gap: 18px;
            }}

            h1 {{
                font-size: 1.8rem;
                margin-bottom: 4px;
                color: var(--primary-dark);
                display: flex;
                align-items: center;
                gap: 8px;
            }}

            .title-pill {{
                background: var(--primary-light);
                color: var(--primary-dark);
                border-radius: 999px;
                padding: 4px 10px;
                font-size: 0.75rem;
                border: 1px solid var(--primary);
            }}

            .title-sub {{
                font-size: 0.9rem;
                color: var(--text-muted);
                margin-top: 4px;
            }}

            form {{
                display: flex;
                flex-direction: column;
                gap: 16px;
            }}

            .field-group {{
                display: flex;
                flex-direction: column;
                gap: 6px;
            }}

            .field-label-row {{
                display: flex;
                justify-content: space-between;
                align-items: center;
                gap: 8px;
            }}

            label {{
                font-weight: 600;
                font-size: 0.9rem;
                color: var(--text-main);
            }}

            .hint {{
                font-size: 0.8rem;
                color: var(--text-muted);
            }}

            .radio-row {{
                display: flex;
                gap: 10px;
                margin-top: 6px;
                flex-wrap: wrap;
            }}

            .radio-option {{
                background: #fff;
                border: 1px solid var(--primary);
                border-radius: 25px;
                padding: 6px 14px;
                cursor: pointer;
                font-size: 0.9rem;
                color: var(--primary-dark);
                display: flex;
                align-items: center;
                gap: 6px;
                transition: 0.2s;
            }}

            .radio-option:hover {{
                background: var(--primary-light);
            }}

            input[type="radio"] {{
                accent-color: var(--primary);
            }}

            textarea,
            input[type="text"] {{
                width: 100%;
                padding: 10px;
                border-radius: 10px;
                border: 1px solid var(--border);
                background: #fff;
                font-size: 0.92rem;
                transition: 0.2s;
            }}

            textarea:focus,
            input[type="text"]:focus {{
                border-color: var(--primary);
                box-shadow: 0 0 0 3px rgba(211, 47, 47, 0.18);
                outline: none;
            }}

            textarea {{
                min-height: 220px;
                resize: vertical;
                font-family: Consolas, Menlo, Monaco, monospace;
            }}

            input[type="file"] {{
                font-size: 0.9rem;
                margin-top: 6px;
            }}

            .text-section,
            .file-section {{
                border: 1px dashed var(--primary-light);
                border-radius: 10px;
                background: #fff7f7;
                padding: 12px;
                flex-direction: column;
                gap: 6px;
            }}

            .text-section {{
                                display: {text_display};
                                flex-direction: column;
                                gap: 6px;
                                border: 2px solid var(--primary);      /* 🔥 adds red outline */
                                border-radius: 10px;
                            }}

            .file-section {{
                display: {file_display};
            }}

            .alert-error {{
                background: var(--error-bg);
                border: 1px solid var(--error-border);
                padding: 10px 12px;
                border-radius: 8px;
                color: var(--error-text);
                font-size: 0.9rem;
                display: flex;
                align-items: center;
                gap: 8px;
            }}

            .btn-primary {{
                background: var(--primary);
                color: white;
                border: none;
                padding: 12px 20px;
                border-radius: 30px;
                font-size: 1rem;
                font-weight: bold;
                cursor: pointer;
                transition: 0.25s;
                box-shadow: 0 6px 14px rgba(211, 47, 47, 0.4);
                display: inline-flex;
                align-items: center;
                gap: 8px;
            }}

            .btn-primary:hover {{
                background: var(--primary-dark);
                transform: translateY(-2px);
            }}

            .btn-primary:active {{
                transform: scale(0.98);
            }}

            .footer-hint {{
                font-size: 0.8rem;
                color: var(--text-muted);
            }}

            @media (max-width: 640px) {{
                .card {{
                    padding: 20px;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="app-shell">
            <div class="card">
                <div>
                    <h1>User Story Generator</h1>
                        <p class="title-sub">
                            Convert conversations or notes into structured Jira-style user stories.
                        </p>
                </div>

                {error_html}

                <form method="post" action="/generate" enctype="multipart/form-data">
                    <div class="field-group">
                        <div class="field-label-row">
                            <label>Choose how you want to provide the conversation.</label>
                        </div>
                        <div class="radio-row">
                            <label class="radio-option">
                                <input type="radio" name="input_mode" value="text" {text_checked} onclick="toggleInput('text')">
                                <span>Text area</span>
                            </label>
                            <label class="radio-option">
                                <input type="radio" name="input_mode" value="file" {file_checked} onclick="toggleInput('file')">
                                <span>Attachment (.txt / .docx)</span>
                            </label>
                        </div>
                    </div>

                    <div class="field-group text-section" id="text-section">
                        <div class="field-label-row">
                            <label for="conversation_text">Conversation / Notes</label>
                            <span class="hint">Paste Teams chat, meeting notes, or any raw text.</span>
                        </div>
                        <textarea id="conversation_text" name="conversation_text">{html.escape(conversation_text)}</textarea>
                    </div>

                    <div class="field-group file-section" id="file-section">
                        <div class="field-label-row">
                            <label for="upload_file">Upload file</label>
                            <span class="hint">Supported: .txt or .docx</span>
                        </div>
                        <input type="file" id="upload_file" name="upload_file" accept=".txt,.docx" />
                        <span class="hint">Ideal for exported meeting notes or requirement documents.</span>
                    </div>

                    <div class="field-group">
                        <div class="field-label-row">
                            <label for="project_key">Project Key</label>
                            <span class="hint">Optional · e.g. CLAIMS, ENGINEERING, FRAUD</span>
                        </div>
                        <input type="text" id="project_key" name="project_key" placeholder="CLAIMS" value="{html.escape(project_key)}" />
                    </div>

                    <div class="field-group">
                        <div class="field-label-row">
                            <span class="footer-hint">
                                The generated Word file groups the content into Jira-style stories with summaries, descriptions, and acceptance criteria.
                            </span>
                            <button type="submit" class="btn-primary">
                                <span>⬇️</span>
                                <span>Generate Word Document</span>
                            </button>
                        </div>
                    </div>
                </form>
            </div>
        </div>

        <script>
            function toggleInput(mode) {{
                var textSection = document.getElementById("text-section");
                var fileSection = document.getElementById("file-section");

                if (mode === "text") {{
                    textSection.style.display = "flex";
                    fileSection.style.display = "none";
                }} else {{
                    textSection.style.display = "none";
                    fileSection.style.display = "flex";
                }}
            }}
        </script>
    </body>
    </html>
    """

    return HTMLResponse(content=html_content)





# ------------- FASTAPI APP ------------- #

app = FastAPI()


# Simple HTML form page
@app.get("/", response_class=HTMLResponse)
async def form_page():
    # No error by default
    return render_form_page()


# Handle form submit and return Word file
@app.post("/generate")
async def generate_word(
    conversation_text: str = Form(""),
    project_key: str = Form(""),
    input_mode: str = Form("text"),
    upload_file: UploadFile | None = File(None),
):
    try:
        # Decide where to get the raw text from
        if input_mode == "file":
            # Must have a file
            if upload_file is None or upload_file.filename == "":
                raise ValueError("Please upload a .txt or .docx file when 'Attachment' is selected.")

            filename = upload_file.filename.lower()

            # Read file bytes
            file_bytes = await upload_file.read()

            if filename.endswith(".txt"):
                # Decode text file
                raw_text = file_bytes.decode("utf-8", errors="ignore")
            elif filename.endswith(".docx"):
                # Read Word document using python-docx
                doc = Document(io.BytesIO(file_bytes))
                raw_text = "\n".join(p.text for p in doc.paragraphs)
            else:
                raise ValueError("Only .txt and .docx files are supported.")
        else:
            # Text mode: use textarea
            raw_text = conversation_text.strip()
            if not raw_text:
                raise ValueError("Please enter some text when 'Text' is selected.")

        # 1) Call OpenAI to generate requirements
        data = generate_requirements_from_text(raw_text)

        # 2) Override project key if user provided one
        if project_key:
            data["project_key"] = project_key

        # 3) Create Word document bytes
        doc_bytes = create_word_bytes_from_requirements(data)

        # 4) Stream as a downloadable file
        filename = f"{data.get('project_key', 'PROJECT')}_requirements.docx"

        return StreamingResponse(
            io.BytesIO(doc_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            },
        )

    except ValueError as ve:
        # Controlled validation errors (missing key, wrong file, empty text, etc.)
        return render_form_page(
            conversation_text=conversation_text,
            project_key=project_key or "",
            input_mode=input_mode,
            error_message=str(ve),
        )

    except Exception as e:
        # Generic fallback
        return render_form_page(
            conversation_text=conversation_text,
            project_key=project_key or "",
            input_mode=input_mode,
            error_message=(
                "Something went wrong while generating the Word file. "
                "Please try again or contact support. "
                f"Details: {str(e)}"
            ),
        )

